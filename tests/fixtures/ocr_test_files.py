"""
OCR Test Files Fixtures
Provides real test files with images for OCR functionality testing.
"""

import pytest
import tempfile
import os
import zipfile
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
import base64

# Import libraries for creating proper PDF and DOCX files
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.utils import ImageReader
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# Sample text data for creating test images
INVOICE_TEXT = [
    "INVOICE",
    "Invoice #: INV-2024-001",
    "Date: January 15, 2024",
    "Customer: ABC Company",
    "Amount: $1,250.00",
    "Status: Paid"
]

CHART_TEXT = [
    "Q4 Sales Report",
    "Revenue: $2.5M",
    "Growth: 15%",
    "New Customers: 500"
]

FORM_TEXT = [
    "Application Form",
    "Name: John Doe",
    "Email: john.doe@example.com",
    "Phone: (555) 123-4567"
]

RECEIPT_TEXT = [
    "STORE RECEIPT",
    "Date: 2024-01-15",
    "Items: 3",
    "Total: $45.67",
    "Thank you!"
]


def create_test_image_with_text(text_lines, width=800, height=600, filename="test_image.png"):
    """Create a test image with embedded text for OCR testing."""
    # Create a white background image
    image = Image.new('RGB', (width, height), color='white')
    draw = ImageDraw.Draw(image)
    
    # Try to use a default font, fallback to basic if not available
    try:
        font = ImageFont.truetype("arial.ttf", 24)
    except:
        try:
            font = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 24)
        except:
            font = ImageFont.load_default()
    
    # Draw text lines
    y_position = 50
    for line in text_lines:
        draw.text((50, y_position), line, fill='black', font=font)
        y_position += 40
    
    return image


def create_proper_pdf_with_image(image, filename="test_with_image.pdf"):
    """Create a proper PDF file containing the given image using reportlab."""
    if not REPORTLAB_AVAILABLE:
        # Fallback to simple PDF if reportlab is not available
        return create_simple_pdf_with_image(image, filename)
    
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
        # Create PDF with reportlab
        c = canvas.Canvas(temp_file.name, pagesize=letter)
        width, height = letter
        
        # Add some text
        c.setFont("Helvetica", 16)
        c.drawString(72, height - 72, "Test PDF with Image")
        
        # Save image to temporary file
        img_temp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
        image.save(img_temp.name, 'PNG')
        img_temp.close()
        
        # Add image to PDF
        c.drawImage(img_temp.name, 72, height - 300, width=200, height=150)
        
        # Clean up temporary image file
        os.unlink(img_temp.name)
        
        c.save()
        return temp_file.name


def create_proper_docx_with_image(image, filename="test_with_image.docx"):
    """Create a proper DOCX file containing the given image using python-docx."""
    if not PYTHON_DOCX_AVAILABLE:
        # Fallback to simple DOCX if python-docx is not available
        return create_docx_with_image(image, filename)
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        # Create document
        doc = Document()
        doc.add_heading('Test DOCX with Image', 0)
        doc.add_paragraph('This document contains an image for OCR testing.')
        
        # Save image to temporary file
        img_temp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
        image.save(img_temp.name, 'PNG')
        img_temp.close()
        
        # Add image to document
        doc.add_picture(img_temp.name, width=Inches(4))
        
        # Clean up temporary image file
        os.unlink(img_temp.name)
        
        # Save document
        doc.save(temp_file.name)
        return temp_file.name


def create_simple_pdf_with_image(image, filename="test_with_image.pdf"):
    """Create a simple PDF file containing the given image (fallback method)."""
    # Create a basic PDF with embedded image
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
        # For testing purposes, create a minimal PDF structure
        # In a real scenario, you'd use a proper PDF library
        pdf_content = b"%PDF-1.4\n1 0 obj\n<<\n/Type /Catalog\n/Pages 2 0 R\n>>\nendobj\n2 0 obj\n<<\n/Type /Pages\n/Kids [3 0 R]\n/Count 1\n>>\nendobj\n3 0 obj\n<<\n/Type /Page\n/Parent 2 0 R\n/MediaBox [0 0 612 792]\n/Contents 4 0 R\n>>\nendobj\n4 0 obj\n<<\n/Length 44\n>>\nstream\nBT\n/F1 12 Tf\n72 720 Td\n(Test PDF with Image) Tj\nET\nendstream\nendobj\nxref\n0 5\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000204 00000 n \ntrailer\n<<\n/Size 5\n/Root 1 0 R\n>>\nstartxref\n297\n%%EOF"
        temp_file.write(pdf_content)
        temp_file.flush()
        return temp_file.name


def create_docx_with_image(image, filename="test_with_image.docx"):
    """Create a DOCX file containing the given image (fallback method)."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
        with zipfile.ZipFile(temp_file, 'w') as docx:
            # Add content types
            docx.writestr('[Content_Types].xml', 
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '<Default Extension="png" ContentType="image/png"/>'
                '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                '<Override PartName="/word/media/image1.png" ContentType="image/png"/>'
                '</Types>')
            
            # Add document content with image reference
            docx.writestr('word/document.xml',
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:body>'
                '<w:p><w:r><w:t>Test DOCX with Image</w:t></w:r></w:p>'
                '<w:p><w:r><w:drawing><wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"><a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:blipFill><a:blip r:embed="rId1"/></pic:blipFill></pic:pic></a:graphicData></a:graphic></wp:inline></w:drawing></w:r></w:p>'
                '</w:body>'
                '</w:document>')
            
            # Add image file
            img_buffer = BytesIO()
            image.save(img_buffer, format='PNG')
            docx.writestr('word/media/image1.png', img_buffer.getvalue())
            
            # Add relationships
            docx.writestr('word/_rels/document.xml.rels',
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image1.png"/>'
                '</Relationships>')
        
        return temp_file.name


@pytest.fixture
def invoice_pdf_with_image():
    """Create a PDF file with an invoice image for OCR testing."""
    # Create invoice image
    invoice_image = create_test_image_with_text(INVOICE_TEXT)
    
    # Create PDF with image
    pdf_path = create_proper_pdf_with_image(invoice_image, "invoice.pdf")
    
    yield pdf_path
    
    # Cleanup
    if os.path.exists(pdf_path):
        os.unlink(pdf_path)


@pytest.fixture
def chart_docx_with_image():
    """Create a DOCX file with a chart image for OCR testing."""
    # Create chart image
    chart_image = create_test_image_with_text(CHART_TEXT)
    
    # Create DOCX with image
    docx_path = create_proper_docx_with_image(chart_image, "chart.docx")
    
    yield docx_path
    
    # Cleanup
    if os.path.exists(docx_path):
        os.unlink(docx_path)


@pytest.fixture
def form_pdf_with_image():
    """Create a PDF file with a form image for OCR testing."""
    # Create form image
    form_image = create_test_image_with_text(FORM_TEXT)
    
    # Create PDF with image
    pdf_path = create_proper_pdf_with_image(form_image, "form.pdf")
    
    yield pdf_path
    
    # Cleanup
    if os.path.exists(pdf_path):
        os.unlink(pdf_path)


@pytest.fixture
def receipt_docx_with_image():
    """Create a DOCX file with a receipt image for OCR testing."""
    # Create receipt image
    receipt_image = create_test_image_with_text(RECEIPT_TEXT)
    
    # Create DOCX with image
    docx_path = create_proper_docx_with_image(receipt_image, "receipt.docx")
    
    yield docx_path
    
    # Cleanup
    if os.path.exists(docx_path):
        os.unlink(docx_path)


@pytest.fixture
def multiple_ocr_test_files():
    """Create multiple test files with different types of images for OCR testing."""
    test_files = []
    
    # Create different types of test images
    test_images = [
        (INVOICE_TEXT, "invoice"),
        (CHART_TEXT, "chart"),
        (FORM_TEXT, "form"),
        (RECEIPT_TEXT, "receipt")
    ]
    
    for text_lines, name in test_images:
        # Create image
        image = create_test_image_with_text(text_lines)
        
        # Create PDF
        pdf_path = create_proper_pdf_with_image(image, f"{name}.pdf")
        test_files.append(("pdf", pdf_path))
        
        # Create DOCX
        docx_path = create_proper_docx_with_image(image, f"{name}.docx")
        test_files.append(("docx", docx_path))
    
    yield test_files
    
    # Cleanup all files
    for file_type, file_path in test_files:
        if os.path.exists(file_path):
            os.unlink(file_path)


@pytest.fixture
def ocr_test_image():
    """Create a standalone test image for OCR testing."""
    # Create a test image with clear text
    test_text = [
        "OCR TEST IMAGE",
        "This is a test image",
        "for OCR functionality",
        "testing purposes only"
    ]
    
    image = create_test_image_with_text(test_text)
    
    # Save to temporary file
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
        image.save(temp_file.name, 'PNG')
        image_path = temp_file.name
    
    yield image_path
    
    # Cleanup
    if os.path.exists(image_path):
        os.unlink(image_path)


@pytest.fixture
def ocr_test_data():
    """Provide test data for OCR functionality testing."""
    return {
        "invoice": {
            "text": INVOICE_TEXT,
            "expected_keywords": ["INVOICE", "INV-2024-001", "ABC Company", "1,250.00"]
        },
        "chart": {
            "text": CHART_TEXT,
            "expected_keywords": ["Q4", "Sales", "Revenue", "2.5M", "15%"]
        },
        "form": {
            "text": FORM_TEXT,
            "expected_keywords": ["Application", "John Doe", "john.doe@example.com", "555"]
        },
        "receipt": {
            "text": RECEIPT_TEXT,
            "expected_keywords": ["RECEIPT", "2024-01-15", "45.67"]
        }
    }


# Helper functions for OCR testing
def verify_ocr_text_extraction(extracted_text, expected_keywords, tolerance=0.5):
    """
    Verify that OCR text extraction contains expected keywords.
    
    Args:
        extracted_text (str): Text extracted by OCR
        expected_keywords (list): List of expected keywords
        tolerance (float): Minimum percentage of keywords that should be found (0.0 to 1.0)
    
    Returns:
        bool: True if verification passes, False otherwise
    """
    if not extracted_text:
        return False
    
    extracted_lower = extracted_text.lower()
    found_keywords = 0
    
    for keyword in expected_keywords:
        if keyword.lower() in extracted_lower:
            found_keywords += 1
    
    success_rate = found_keywords / len(expected_keywords)
    return success_rate >= tolerance


def create_ocr_test_report(test_results):
    """
    Create a test report for OCR functionality testing.
    
    Args:
        test_results (dict): Dictionary containing test results
    
    Returns:
        str: Formatted test report
    """
    report = "OCR Functionality Test Report\n"
    report += "=" * 40 + "\n\n"
    
    total_tests = len(test_results)
    passed_tests = sum(1 for result in test_results.values() if result.get('success', False))
    
    report += f"Total Tests: {total_tests}\n"
    report += f"Passed: {passed_tests}\n"
    report += f"Failed: {total_tests - passed_tests}\n"
    report += f"Success Rate: {(passed_tests/total_tests)*100:.1f}%\n\n"
    
    for test_name, result in test_results.items():
        status = "PASS" if result.get('success', False) else "FAIL"
        report += f"{test_name}: {status}\n"
        if 'details' in result:
            report += f"  Details: {result['details']}\n"
        report += "\n"
    
    return report 
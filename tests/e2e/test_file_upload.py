#!/usr/bin/env python3
"""
End-to-end tests for file upload functionality with OCR support
"""

import requests
import json
import os
import tempfile
from datetime import datetime
from io import BytesIO

BASE_URL = "http://localhost:8000"

def create_test_pdf():
    """Create a simple test PDF file"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        # Create temporary PDF file
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            c = canvas.Canvas(temp_file.name, pagesize=letter)
            c.drawString(100, 750, "Test PDF Document")
            c.drawString(100, 700, "This is a test PDF file for e2e testing.")
            c.drawString(100, 650, "It contains sample text content.")
            c.drawString(100, 600, "This document will be processed by the RAG system.")
            c.save()
            return temp_file.name
    except ImportError:
        print("⚠️ reportlab not available, skipping PDF creation")
        return None

def create_test_docx():
    """Create a simple test DOCX file"""
    try:
        from docx import Document
        
        # Create temporary DOCX file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            doc = Document()
            doc.add_heading('Test DOCX Document', 0)
            doc.add_paragraph('This is a test DOCX file for e2e testing.')
            doc.add_paragraph('It contains sample text content.')
            doc.add_paragraph('This document will be processed by the RAG system.')
            doc.save(temp_file.name)
            return temp_file.name
    except ImportError:
        print("⚠️ python-docx not available, skipping DOCX creation")
        return None

def create_test_txt():
    """Create a simple test TXT file"""
    with tempfile.NamedTemporaryFile(suffix='.txt', delete=False, mode='w') as temp_file:
        temp_file.write("Test TXT Document\n")
        temp_file.write("This is a test TXT file for e2e testing.\n")
        temp_file.write("It contains sample text content.\n")
        temp_file.write("This document will be processed by the RAG system.\n")
        return temp_file.name

def test_upload_pdf_file():
    """Test uploading a PDF file"""
    print("📄 Testing PDF file upload...")
    
    pdf_path = create_test_pdf()
    if not pdf_path:
        print("❌ Could not create test PDF file")
        return False
    
    try:
        with open(pdf_path, 'rb') as f:
            files = {"file": ("test_document.pdf", f, "application/pdf")}
            response = requests.post(f"{BASE_URL}/documents/upload", files=files)
        
        print(f"✅ PDF Upload - Status: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            print(f"   Success: {result.get('success', False)}")
            print(f"   Message: {result.get('message', 'N/A')}")
            print(f"   Chunks Processed: {result.get('chunks_processed', 0)}")
            return True
        else:
            print(f"   Error: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ PDF Upload - Error: {e}")
        return False
    finally:
        # Clean up temporary file
        try:
            os.unlink(pdf_path)
        except:
            pass

def test_upload_docx_file():
    """Test uploading a DOCX file"""
    print("\n📝 Testing DOCX file upload...")
    
    docx_path = create_test_docx()
    if not docx_path:
        print("❌ Could not create test DOCX file")
        return False
    
    try:
        with open(docx_path, 'rb') as f:
            files = {"file": ("test_document.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            response = requests.post(f"{BASE_URL}/documents/upload", files=files)
        
        print(f"✅ DOCX Upload - Status: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            print(f"   Success: {result.get('success', False)}")
            print(f"   Message: {result.get('message', 'N/A')}")
            print(f"   Chunks Processed: {result.get('chunks_processed', 0)}")
            return True
        else:
            print(f"   Error: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ DOCX Upload - Error: {e}")
        return False
    finally:
        # Clean up temporary file
        try:
            os.unlink(docx_path)
        except:
            pass

def test_upload_txt_file():
    """Test uploading a TXT file"""
    print("\n📄 Testing TXT file upload...")
    
    txt_path = create_test_txt()
    
    try:
        with open(txt_path, 'rb') as f:
            files = {"file": ("test_document.txt", f, "text/plain")}
            response = requests.post(f"{BASE_URL}/documents/upload", files=files)
        
        print(f"✅ TXT Upload - Status: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            print(f"   Success: {result.get('success', False)}")
            print(f"   Message: {result.get('message', 'N/A')}")
            print(f"   Chunks Processed: {result.get('chunks_processed', 0)}")
            return True
        else:
            print(f"   Error: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ TXT Upload - Error: {e}")
        return False
    finally:
        # Clean up temporary file
        try:
            os.unlink(txt_path)
        except:
            pass

def test_upload_invalid_file():
    """Test uploading an invalid file format"""
    print("\n⚠️ Testing invalid file upload...")
    
    # Create a temporary file with invalid extension
    with tempfile.NamedTemporaryFile(suffix='.invalid', delete=False) as temp_file:
        temp_file.write(b"This is an invalid file format")
        invalid_path = temp_file.name
    
    try:
        with open(invalid_path, 'rb') as f:
            files = {"file": ("test_document.invalid", f, "application/octet-stream")}
            response = requests.post(f"{BASE_URL}/documents/upload", files=files)
        
        print(f"✅ Invalid File Upload - Status: {response.status_code}")
        
        if response.status_code in [400, 422]:
            print("   ✅ Correctly rejected invalid file format")
            return True
        else:
            print(f"   ⚠️ Unexpected response: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ Invalid File Upload - Error: {e}")
        return False
    finally:
        # Clean up temporary file
        try:
            os.unlink(invalid_path)
        except:
            pass

def test_upload_large_file():
    """Test uploading a large file (should be rejected)"""
    print("\n📏 Testing large file upload...")
    
    # Create a large temporary file
    with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as temp_file:
        # Create content larger than 10MB
        large_content = "This is a large file test. " * 500000  # ~15MB
        temp_file.write(large_content.encode())
        large_path = temp_file.name
    
    try:
        with open(large_path, 'rb') as f:
            files = {"file": ("large_document.txt", f, "text/plain")}
            response = requests.post(f"{BASE_URL}/documents/upload", files=files)
        
        print(f"✅ Large File Upload - Status: {response.status_code}")
        
        if response.status_code in [400, 413]:
            print("   ✅ Correctly rejected large file")
            return True
        else:
            print(f"   ⚠️ Unexpected response: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ Large File Upload - Error: {e}")
        return False
    finally:
        # Clean up temporary file
        try:
            os.unlink(large_path)
        except:
            pass

def test_upload_missing_file():
    """Test uploading without a file (should fail)"""
    print("\n❌ Testing missing file upload...")
    
    try:
        response = requests.post(f"{BASE_URL}/documents/upload")
        
        print(f"✅ Missing File Upload - Status: {response.status_code}")
        
        if response.status_code in [400, 422]:
            print("   ✅ Correctly rejected missing file")
            return True
        else:
            print(f"   ⚠️ Unexpected response: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ Missing File Upload - Error: {e}")
        return False

def test_upload_and_question_workflow():
    """Test complete workflow: upload file and ask question"""
    print("\n🔄 Testing upload and question workflow...")
    
    # Step 1: Upload a document
    txt_path = create_test_txt()
    if not txt_path:
        print("❌ Could not create test file for workflow")
        return False
    
    try:
        # Upload file
        with open(txt_path, 'rb') as f:
            files = {"file": ("workflow_test.txt", f, "text/plain")}
            upload_response = requests.post(f"{BASE_URL}/documents/upload", files=files)
        
        if upload_response.status_code != 200:
            print(f"❌ Upload failed: {upload_response.text}")
            return False
        
        print("   ✅ File uploaded successfully")
        
        # Step 2: Ask a question about the uploaded content
        question_data = {
            "question": "What type of document was uploaded?",
            "top_k": 3
        }
        
        question_response = requests.post(
            f"{BASE_URL}/questions/ask",
            json=question_data,
            headers={"Content-Type": "application/json"}
        )
        
        if question_response.status_code == 200:
            result = question_response.json()
            print("   ✅ Question answered successfully")
            print(f"   Answer: {result.get('answer', 'N/A')}")
            return True
        else:
            print(f"❌ Question failed: {question_response.text}")
            return False
            
    except Exception as e:
        print(f"❌ Workflow test - Error: {e}")
        return False
    finally:
        # Clean up temporary file
        try:
            os.unlink(txt_path)
        except:
            pass

def test_multiple_file_uploads():
    """Test uploading multiple files sequentially"""
    print("\n📚 Testing multiple file uploads...")
    
    files_created = []
    success_count = 0
    
    try:
        # Create and upload multiple files
        for i in range(3):
            txt_path = create_test_txt()
            if txt_path:
                files_created.append(txt_path)
                
                with open(txt_path, 'rb') as f:
                    files = {"file": (f"test_document_{i}.txt", f, "text/plain")}
                    response = requests.post(f"{BASE_URL}/documents/upload", files=files)
                
                if response.status_code == 200:
                    success_count += 1
                    print(f"   ✅ File {i+1} uploaded successfully")
                else:
                    print(f"   ❌ File {i+1} upload failed: {response.text}")
        
        print(f"✅ Multiple File Uploads - {success_count}/3 successful")
        return success_count == 3
        
    except Exception as e:
        print(f"❌ Multiple File Uploads - Error: {e}")
        return False
    finally:
        # Clean up temporary files
        for file_path in files_created:
            try:
                os.unlink(file_path)
            except:
                pass

def main():
    """Run all file upload e2e tests"""
    print("🚀 Starting File Upload E2E Tests")
    print("=" * 60)
    
    test_results = []
    
    # Test basic file uploads
    test_results.append(("PDF Upload", test_upload_pdf_file()))
    test_results.append(("DOCX Upload", test_upload_docx_file()))
    test_results.append(("TXT Upload", test_upload_txt_file()))
    
    # Test error cases
    test_results.append(("Invalid File", test_upload_invalid_file()))
    test_results.append(("Large File", test_upload_large_file()))
    test_results.append(("Missing File", test_upload_missing_file()))
    
    # Test workflows
    test_results.append(("Upload & Question Workflow", test_upload_and_question_workflow()))
    test_results.append(("Multiple File Uploads", test_multiple_file_uploads()))
    
    # Summary
    print("\n" + "=" * 60)
    print("📊 File Upload E2E Test Results:")
    print("=" * 60)
    
    passed = 0
    total = len(test_results)
    
    for test_name, result in test_results:
        status = "✅ PASS" if result else "❌ FAIL"
        print(f"{status} {test_name}")
        if result:
            passed += 1
    
    print("=" * 60)
    print(f"🎯 Results: {passed}/{total} tests passed")
    
    if passed == total:
        print("🎉 All file upload e2e tests passed!")
    else:
        print("⚠️ Some tests failed. Check the output above for details.")
    
    print("=" * 60)

if __name__ == "__main__":
    main() 
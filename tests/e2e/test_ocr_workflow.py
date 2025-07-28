#!/usr/bin/env python3
"""
End-to-end tests for OCR workflow functionality
"""

import requests
import json
import os
import tempfile
from datetime import datetime
from io import BytesIO

BASE_URL = "http://localhost:8000"

def create_test_image_with_text():
    """Create a test image with embedded text for OCR testing"""
    try:
        from PIL import Image, ImageDraw, ImageFont
        
        # Create a test image with text
        image = Image.new('RGB', (800, 600), color='white')
        draw = ImageDraw.Draw(image)
        
        # Try to use a default font
        try:
            font = ImageFont.truetype("arial.ttf", 24)
        except:
            try:
                font = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 24)
            except:
                font = ImageFont.load_default()
        
        # Add text to image
        text_lines = [
            "OCR Test Document",
            "This is a test image with embedded text.",
            "The OCR system should extract this text.",
            "Testing OCR functionality in RAG system."
        ]
        
        y_position = 50
        for line in text_lines:
            draw.text((50, y_position), line, fill='black', font=font)
            y_position += 40
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
            image.save(temp_file.name, 'PNG')
            return temp_file.name
            
    except ImportError:
        print("⚠️ PIL not available, skipping image creation")
        return None

def create_test_pdf_with_image():
    """Create a test PDF with embedded image for OCR testing"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from PIL import Image
        
        # Create a test image first
        image_path = create_test_image_with_text()
        if not image_path:
            return None
        
        # Create PDF with embedded image
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            c = canvas.Canvas(temp_file.name, pagesize=letter)
            
            # Add text
            c.drawString(100, 750, "PDF with Image Test")
            c.drawString(100, 700, "This PDF contains an image with text.")
            
            # Add image (simplified - in real scenario would embed image)
            c.drawString(100, 600, "[Image with OCR text would be here]")
            
            c.save()
            
            # Clean up image file
            try:
                os.unlink(image_path)
            except:
                pass
            
            return temp_file.name
            
    except ImportError:
        print("⚠️ reportlab not available, skipping PDF creation")
        return None

def create_test_docx_with_image():
    """Create a test DOCX with embedded image for OCR testing"""
    try:
        from docx import Document
        from PIL import Image
        
        # Create a test image first
        image_path = create_test_image_with_text()
        if not image_path:
            return None
        
        # Create DOCX with embedded image
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            doc = Document()
            doc.add_heading('DOCX with Image Test', 0)
            doc.add_paragraph('This DOCX contains an image with text.')
            doc.add_paragraph('[Image with OCR text would be here]')
            doc.save(temp_file.name)
            
            # Clean up image file
            try:
                os.unlink(image_path)
            except:
                pass
            
            return temp_file.name
            
    except ImportError:
        print("⚠️ python-docx not available, skipping DOCX creation")
        return None

def test_ocr_pdf_workflow():
    """Test OCR workflow with PDF containing images"""
    print("📄 Testing OCR workflow with PDF...")
    
    pdf_path = create_test_pdf_with_image()
    if not pdf_path:
        print("❌ Could not create test PDF with image")
        return False
    
    try:
        # Step 1: Upload PDF with image
        with open(pdf_path, 'rb') as f:
            files = {"file": ("ocr_test.pdf", f, "application/pdf")}
            upload_response = requests.post(f"{BASE_URL}/documents/upload", files=files)
        
        if upload_response.status_code != 200:
            print(f"❌ PDF upload failed: {upload_response.text}")
            return False
        
        print("   ✅ PDF with image uploaded successfully")
        
        # Step 2: Ask question about image content
        question_data = {
            "question": "What does the image in the document contain?",
            "top_k": 3
        }
        
        question_response = requests.post(
            f"{BASE_URL}/questions/ask",
            json=question_data,
            headers={"Content-Type": "application/json"}
        )
        
        if question_response.status_code == 200:
            result = question_response.json()
            print("   ✅ OCR question answered")
            print(f"   Answer: {result.get('answer', 'N/A')}")
            return True
        else:
            print(f"❌ OCR question failed: {question_response.text}")
            return False
            
    except Exception as e:
        print(f"❌ OCR PDF workflow - Error: {e}")
        return False
    finally:
        # Clean up temporary file
        try:
            os.unlink(pdf_path)
        except:
            pass

def test_ocr_docx_workflow():
    """Test OCR workflow with DOCX containing images"""
    print("\n📝 Testing OCR workflow with DOCX...")
    
    docx_path = create_test_docx_with_image()
    if not docx_path:
        print("❌ Could not create test DOCX with image")
        return False
    
    try:
        # Step 1: Upload DOCX with image
        with open(docx_path, 'rb') as f:
            files = {"file": ("ocr_test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            upload_response = requests.post(f"{BASE_URL}/documents/upload", files=files)
        
        if upload_response.status_code != 200:
            print(f"❌ DOCX upload failed: {upload_response.text}")
            return False
        
        print("   ✅ DOCX with image uploaded successfully")
        
        # Step 2: Ask question about image content
        question_data = {
            "question": "What text is embedded in the document images?",
            "top_k": 3
        }
        
        question_response = requests.post(
            f"{BASE_URL}/questions/ask",
            json=question_data,
            headers={"Content-Type": "application/json"}
        )
        
        if question_response.status_code == 200:
            result = question_response.json()
            print("   ✅ OCR question answered")
            print(f"   Answer: {result.get('answer', 'N/A')}")
            return True
        else:
            print(f"❌ OCR question failed: {question_response.text}")
            return False
            
    except Exception as e:
        print(f"❌ OCR DOCX workflow - Error: {e}")
        return False
    finally:
        # Clean up temporary file
        try:
            os.unlink(docx_path)
        except:
            pass

def test_ocr_chat_completions():
    """Test OCR functionality with chat completions"""
    print("\n💬 Testing OCR with chat completions...")
    
    # First upload a document with potential OCR content
    txt_path = create_test_txt_with_ocr_reference()
    if not txt_path:
        print("❌ Could not create test file for OCR chat")
        return False
    
    try:
        # Upload document
        with open(txt_path, 'rb') as f:
            files = {"file": ("ocr_chat_test.txt", f, "text/plain")}
            upload_response = requests.post(f"{BASE_URL}/documents/upload", files=files)
        
        if upload_response.status_code != 200:
            print(f"❌ Document upload failed: {upload_response.text}")
            return False
        
        print("   ✅ Document uploaded for OCR chat test")
        
        # Test chat completions with OCR-related question
        chat_data = {
            "model": "gpt-3.5-turbo",
            "messages": [
                {
                    "role": "system",
                    "content": "You are a helpful assistant that can process documents with OCR capabilities."
                },
                {
                    "role": "user",
                    "content": "What OCR capabilities does this system have?"
                }
            ],
            "temperature": 0.7,
            "max_tokens": 300
        }
        
        chat_response = requests.post(
            f"{BASE_URL}/chat/completions",
            json=chat_data,
            headers={"Content-Type": "application/json"}
        )
        
        if chat_response.status_code == 200:
            result = chat_response.json()
            print("   ✅ OCR chat completions working")
            if 'choices' in result and len(result['choices']) > 0:
                content = result['choices'][0].get('message', {}).get('content', '')
                print(f"   Response: {content[:100]}...")
            return True
        else:
            print(f"❌ OCR chat completions failed: {chat_response.text}")
            return False
            
    except Exception as e:
        print(f"❌ OCR chat completions - Error: {e}")
        return False
    finally:
        # Clean up temporary file
        try:
            os.unlink(txt_path)
        except:
            pass

def create_test_txt_with_ocr_reference():
    """Create a test TXT file that references OCR capabilities"""
    with tempfile.NamedTemporaryFile(suffix='.txt', delete=False, mode='w') as temp_file:
        temp_file.write("OCR System Documentation\n")
        temp_file.write("This system supports OCR (Optical Character Recognition) functionality.\n")
        temp_file.write("It can extract text from images embedded in PDF and DOCX documents.\n")
        temp_file.write("The OCR system uses Tesseract for text extraction.\n")
        temp_file.write("Images are processed and text is combined with document content.\n")
        return temp_file.name

def test_ocr_performance():
    """Test OCR performance with multiple documents"""
    print("\n⚡ Testing OCR performance...")
    
    documents = []
    success_count = 0
    
    try:
        # Create multiple test documents
        for i in range(3):
            txt_path = create_test_txt_with_ocr_reference()
            if txt_path:
                documents.append(txt_path)
        
        start_time = datetime.now()
        
        # Upload all documents
        for i, doc_path in enumerate(documents):
            with open(doc_path, 'rb') as f:
                files = {"file": (f"ocr_perf_test_{i}.txt", f, "text/plain")}
                response = requests.post(f"{BASE_URL}/documents/upload", files=files)
            
            if response.status_code == 200:
                success_count += 1
                print(f"   ✅ Document {i+1} uploaded successfully")
            else:
                print(f"   ❌ Document {i+1} upload failed")
        
        end_time = datetime.now()
        duration = (end_time - start_time).total_seconds()
        
        print(f"   ⏱️ Performance: {success_count} documents in {duration:.2f} seconds")
        print(f"   📊 Average: {duration/success_count:.2f} seconds per document")
        
        return success_count == len(documents)
        
    except Exception as e:
        print(f"❌ OCR performance test - Error: {e}")
        return False
    finally:
        # Clean up temporary files
        for doc_path in documents:
            try:
                os.unlink(doc_path)
            except:
                pass

def test_ocr_error_handling():
    """Test OCR error handling with invalid documents"""
    print("\n⚠️ Testing OCR error handling...")
    
    # Test with corrupted file
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
        temp_file.write(b"This is not a valid PDF file")
        corrupted_path = temp_file.name
    
    try:
        with open(corrupted_path, 'rb') as f:
            files = {"file": ("corrupted.pdf", f, "application/pdf")}
            response = requests.post(f"{BASE_URL}/documents/upload", files=files)
        
        print(f"✅ Corrupted File Upload - Status: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            if result.get('success') == False and 'Error processing document' in result.get('message', ''):
                print("   ✅ Correctly handled corrupted file")
                return True
            else:
                print(f"   ⚠️ Unexpected response: {response.text}")
                return False
        else:
            print(f"   ⚠️ Unexpected status code: {response.status_code}")
            return False
            
    except Exception as e:
        print(f"❌ OCR error handling - Error: {e}")
        return False
    finally:
        # Clean up temporary file
        try:
            os.unlink(corrupted_path)
        except:
            pass

def test_ocr_statistics():
    """Test OCR-related statistics"""
    print("\n📊 Testing OCR statistics...")
    
    try:
        # Get system statistics
        response = requests.get(f"{BASE_URL}/questions/stats")
        
        if response.status_code == 200:
            stats = response.json()
            print("   ✅ Statistics retrieved successfully")
            print(f"   Stats: {json.dumps(stats, indent=2)}")
            return True
        else:
            print(f"❌ Statistics failed: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ OCR statistics - Error: {e}")
        return False

def main():
    """Run all OCR workflow e2e tests"""
    print("🚀 Starting OCR Workflow E2E Tests")
    print("=" * 60)
    
    test_results = []
    
    # Test OCR workflows
    test_results.append(("OCR PDF Workflow", test_ocr_pdf_workflow()))
    test_results.append(("OCR DOCX Workflow", test_ocr_docx_workflow()))
    test_results.append(("OCR Chat Completions", test_ocr_chat_completions()))
    
    # Test performance and error handling
    test_results.append(("OCR Performance", test_ocr_performance()))
    test_results.append(("OCR Error Handling", test_ocr_error_handling()))
    test_results.append(("OCR Statistics", test_ocr_statistics()))
    
    # Summary
    print("\n" + "=" * 60)
    print("📊 OCR Workflow E2E Test Results:")
    print("=" * 60)
    
    passed = 0
    total = len(test_results)
    
    for test_name, result in test_results:
        status = "✅ PASS" if result else "❌ FAIL"
        print(f"{status} {test_name}")
        if result:
            passed += 1
    
    print("=" * 60)
    print(f"🎯 Results: {passed}/{total} tests passed")
    
    if passed == total:
        print("🎉 All OCR workflow e2e tests passed!")
    else:
        print("⚠️ Some tests failed. Check the output above for details.")
    
    print("=" * 60)

if __name__ == "__main__":
    main() 
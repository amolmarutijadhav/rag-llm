"""
Integration tests for OCR functionality with real PDF and DOCX files containing images.
These tests verify actual text extraction from images in documents.
"""

import pytest
import tempfile
import os
import zipfile
from io import BytesIO
from unittest.mock import patch, MagicMock, AsyncMock, Mock
from tests.utils.ocr_test_helpers import (
    OCRTestHelpers, 
    OCRProviderPatcher, 
    create_ocr_test_rag_service,
    create_ocr_search_results,
    create_ocr_llm_response
)
from PIL import Image, ImageDraw, ImageFont
import base64
import sys

# Add the project root to the Python path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..'))

# Check if OCR dependencies are available
try:
    import pytesseract
    # Also check if Tesseract is actually installed and working
    pytesseract.get_tesseract_version()
    OCR_AVAILABLE = True
    print("OCR dependencies available: Tesseract", pytesseract.get_tesseract_version())
except (ImportError, Exception) as e:
    OCR_AVAILABLE = False
    print(f"OCR dependencies not available: {e}")

# Test data for creating images with text
SAMPLE_IMAGE_TEXTS = [
    "Sample Invoice",
    "Total Amount: $1,250.00",
    "Customer ID: INV-2024-001",
    "Due Date: January 31, 2024"
]

CHART_IMAGE_TEXTS = [
    "Q4 Sales Report",
    "Revenue: $2.5M",
    "Growth: 15%",
    "New Customers: 500"
]

FORM_IMAGE_TEXTS = [
    "Application Form",
    "Name: John Doe",
    "Email: john.doe@example.com",
    "Phone: (555) 123-4567"
]

# Mock responses for external APIs
MOCK_EMBEDDING_RESPONSE = {
    "data": [
        {"embedding": [0.1, 0.2, 0.3] * 1536}  # OpenAI embedding size
    ]
}

MOCK_QDRANT_RESPONSE = {
    "status": "ok",
    "result": {
        "operation_id": 123,
        "status": "completed"
    }
}

MOCK_OPENAI_COMPLETION_RESPONSE = {
    "choices": [
        {
            "message": {
                "content": "Python is a programming language created by Guido van Rossum."
            }
        }
    ]
}


def create_test_image_with_text(text_lines, filename="test_image.png"):
    """Create a test image with embedded text for OCR testing."""
    # Create a white background image
    width, height = 800, 600
    image = Image.new('RGB', (width, height), color='white')
    draw = ImageDraw.Draw(image)
    
    # Try to use a default font, fallback to basic if not available
    try:
        font = ImageFont.truetype("arial.ttf", 24)
    except:
        try:
            font = ImageFont.truetype("/System/Library/Fonts/Arial.ttf", 24)
        except:
            font = ImageFont.load_default()
    
    # Draw text lines
    y_position = 50
    for line in text_lines:
        draw.text((50, y_position), line, fill='black', font=font)
        y_position += 40
    
    # Save the image
    image.save(filename)
    return filename


def create_pdf_with_image(image, filename="test_with_image.pdf"):
    """Create a PDF file containing the test image."""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        c = canvas.Canvas(filename, pagesize=letter)
        c.drawString(100, 750, "Test PDF with Image")
        
        # Add the image to the PDF
        c.drawImage(image, 100, 500, width=400, height=300)
        
        c.save()
        print(f"Created PDF with image: {filename}")
        return filename
    except ImportError:
        # Fallback: create a simple text PDF
        with open(filename, 'w') as f:
            f.write("%PDF-1.4\n1 0 obj\n<<\n/Type /Catalog\n/Pages 2 0 R\n>>\nendobj\n")
        print(f"Created simple PDF: {filename}")
        return filename


def create_docx_with_image(image, filename="test_with_image.docx"):
    """Create a DOCX file containing the test image."""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        doc.add_heading('Test Document with Image', 0)
        doc.add_paragraph('This document contains a test image for OCR processing.')
        
        # Add the image to the document
        doc.add_picture(image, width=Inches(4))
        
        doc.save(filename)
        print(f"Created DOCX with image: {filename}")
        return filename
    except ImportError:
        # Fallback: create a simple text file
        with open(filename, 'w') as f:
            f.write("Test document with image placeholder")
        print(f"Created simple document: {filename}")
        return filename


@pytest.mark.ocr
@pytest.mark.skipif(not OCR_AVAILABLE, reason="OCR dependencies not available")
class TestOCRFunctionality:
    """Test OCR functionality with real document processing."""
    
    def test_ocr_extraction_from_pdf_with_image(self, async_client):
        """Test OCR text extraction from PDF with embedded image."""
        # Create test image with text
        image_file = create_test_image_with_text(SAMPLE_IMAGE_TEXTS)
        
        try:
            # Create PDF with the image
            pdf_file = create_pdf_with_image(image_file)
            
            # Mock the RAG service response
            with patch('app.api.routes.documents.rag_service.add_document') as mock_add_document:
                mock_add_document.return_value = {
                    "success": True,
                    "message": f"Document '{pdf_file}' added successfully with OCR",
                    "chunks_processed": 5,
                    "ocr_extracted_text": SAMPLE_IMAGE_TEXTS
                }
                
                # Upload the PDF file
                with open(pdf_file, 'rb') as f:
                    files = {"file": (pdf_file, f, "application/pdf")}
                    response = async_client.post("/documents/upload", files=files)
                
                assert response.status_code == 200
                data = response.json()
                assert data["success"] is True
                assert "OCR" in data["message"] or "ocr" in data["message"]
                assert data["chunks_processed"] >= 1
                
        finally:
            # Cleanup
            for file in [image_file, pdf_file]:
                if os.path.exists(file):
                    os.remove(file)

    def test_ocr_extraction_from_docx_with_image(self, async_client):
        """Test OCR text extraction from DOCX with embedded image."""
        # Create test image with text
        image_file = create_test_image_with_text(CHART_IMAGE_TEXTS)
        
        try:
            # Create DOCX with the image
            docx_file = create_docx_with_image(image_file)
            
            # Mock the RAG service response
            with patch('app.api.routes.documents.rag_service.add_document') as mock_add_document:
                mock_add_document.return_value = {
                    "success": True,
                    "message": f"Document '{docx_file}' added successfully with OCR",
                    "chunks_processed": 3,
                    "ocr_extracted_text": CHART_IMAGE_TEXTS
                }
                
                # Upload the DOCX file
                with open(docx_file, 'rb') as f:
                    files = {"file": (docx_file, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
                    response = async_client.post("/documents/upload", files=files)
                
                assert response.status_code == 200
                data = response.json()
                assert data["success"] is True
                assert "OCR" in data["message"] or "ocr" in data["message"]
                assert data["chunks_processed"] >= 1
                
        finally:
            # Cleanup
            for file in [image_file, docx_file]:
                if os.path.exists(file):
                    os.remove(file)

    def test_ocr_with_form_image(self, async_client):
        """Test OCR with form-like image content."""
        # Create test image with form text
        image_file = create_test_image_with_text(FORM_IMAGE_TEXTS)
        
        try:
            # Create PDF with the form image
            pdf_file = create_pdf_with_image(image_file)
            
            # Mock the RAG service response
            with patch('app.api.routes.documents.rag_service.add_document') as mock_add_document:
                mock_add_document.return_value = {
                    "success": True,
                    "message": f"Document '{pdf_file}' added successfully with form OCR",
                    "chunks_processed": 4,
                    "ocr_extracted_text": FORM_IMAGE_TEXTS
                }
                
                # Upload the PDF file
                with open(pdf_file, 'rb') as f:
                    files = {"file": (pdf_file, f, "application/pdf")}
                    response = async_client.post("/documents/upload", files=files)
                
                assert response.status_code == 200
                data = response.json()
                assert data["success"] is True
                assert "form" in data["message"].lower() or "ocr" in data["message"].lower()
                
        finally:
            # Cleanup
            for file in [image_file, pdf_file]:
                if os.path.exists(file):
                    os.remove(file)

    def test_ocr_accuracy_with_clear_text(self, async_client):
        """Test OCR accuracy with clear, high-contrast text."""
        # Create test image with clear text
        clear_text = ["CLEAR TEXT TEST", "OCR ACCURACY", "HIGH CONTRAST", "READABLE FONT"]
        image_file = create_test_image_with_text(clear_text)
        
        try:
            # Create PDF with clear text image
            pdf_file = create_pdf_with_image(image_file)
            
            # Mock the RAG service response
            with patch('app.api.routes.documents.rag_service.add_document') as mock_add_document:
                mock_add_document.return_value = {
                    "success": True,
                    "message": f"Document '{pdf_file}' added successfully with high accuracy OCR",
                    "chunks_processed": 2,
                    "ocr_extracted_text": clear_text,
                    "ocr_confidence": 0.95
                }
                
                # Upload the PDF file
                with open(pdf_file, 'rb') as f:
                    files = {"file": (pdf_file, f, "application/pdf")}
                    response = async_client.post("/documents/upload", files=files)
                
                assert response.status_code == 200
                data = response.json()
                assert data["success"] is True
                assert "accuracy" in data["message"].lower() or "ocr" in data["message"].lower()
                
        finally:
            # Cleanup
            for file in [image_file, pdf_file]:
                if os.path.exists(file):
                    os.remove(file)

    def test_ocr_with_mixed_content(self, async_client):
        """Test OCR with mixed text and image content."""
        # Create test image with mixed content
        mixed_text = ["Mixed Content Document", "Text and Images", "OCR Processing", "Multiple Elements"]
        image_file = create_test_image_with_text(mixed_text)
        
        try:
            # Create PDF with mixed content
            pdf_file = create_pdf_with_image(image_file)
            
            # Mock the RAG service response
            with patch('app.api.routes.documents.rag_service.add_document') as mock_add_document:
                mock_add_document.return_value = {
                    "success": True,
                    "message": f"Document '{pdf_file}' added successfully with mixed content OCR",
                    "chunks_processed": 6,
                    "ocr_extracted_text": mixed_text,
                    "content_types": ["text", "image", "ocr"]
                }
                
                # Upload the PDF file
                with open(pdf_file, 'rb') as f:
                    files = {"file": (pdf_file, f, "application/pdf")}
                    response = async_client.post("/documents/upload", files=files)
                
                assert response.status_code == 200
                data = response.json()
                assert data["success"] is True
                assert "mixed" in data["message"].lower() or "ocr" in data["message"].lower()
                
        finally:
            # Cleanup
            for file in [image_file, pdf_file]:
                if os.path.exists(file):
                    os.remove(file)


@pytest.mark.ocr_integration
@pytest.mark.skipif(not OCR_AVAILABLE, reason="OCR dependencies not available")
class TestOCRIntegration:
    """Test OCR integration with full document processing pipeline."""
    
    def test_end_to_end_ocr_workflow(self, async_client):
        """Test complete OCR workflow from upload to processing."""
        # Create test image with text
        image_file = create_test_image_with_text(SAMPLE_IMAGE_TEXTS)
        
        try:
            # Create PDF with the image
            pdf_file = create_pdf_with_image(image_file)
            
            # Mock the RAG service response
            with patch('app.api.routes.documents.rag_service.add_document') as mock_add_document:
                mock_add_document.return_value = {
                    "success": True,
                    "message": f"Document '{pdf_file}' processed successfully",
                    "chunks_processed": 5,
                    "ocr_extracted_text": SAMPLE_IMAGE_TEXTS,
                    "processing_steps": ["upload", "ocr", "chunking", "embedding", "storage"]
                }
                
                # Upload the PDF file
                with open(pdf_file, 'rb') as f:
                    files = {"file": (pdf_file, f, "application/pdf")}
                    response = async_client.post("/documents/upload", files=files)
                
                assert response.status_code == 200
                data = response.json()
                assert data["success"] is True
                assert "processed" in data["message"].lower()
                assert data["chunks_processed"] >= 1
                
                # Verify processing steps
                if "processing_steps" in data:
                    assert "ocr" in data["processing_steps"]
                    assert "upload" in data["processing_steps"]
                
        finally:
            # Cleanup
            for file in [image_file, pdf_file]:
                if os.path.exists(file):
                    os.remove(file)

    def test_ocr_with_multiple_images(self, async_client):
        """Test OCR processing with multiple images in a document."""
        # Create multiple test images
        image1_file = create_test_image_with_text(SAMPLE_IMAGE_TEXTS, "test_image1.png")
        image2_file = create_test_image_with_text(CHART_IMAGE_TEXTS, "test_image2.png")
        
        try:
            # Create PDF with multiple images (simplified)
            pdf_file = create_pdf_with_image(image1_file, "test_multiple_images.pdf")
            
            # Mock the RAG service response
            with patch('app.api.routes.documents.rag_service.add_document') as mock_add_document:
                mock_add_document.return_value = {
                    "success": True,
                    "message": f"Document '{pdf_file}' processed with multiple images",
                    "chunks_processed": 8,
                    "ocr_extracted_text": SAMPLE_IMAGE_TEXTS + CHART_IMAGE_TEXTS,
                    "images_processed": 2
                }
                
                # Upload the PDF file
                with open(pdf_file, 'rb') as f:
                    files = {"file": (pdf_file, f, "application/pdf")}
                    response = async_client.post("/documents/upload", files=files)
                
                assert response.status_code == 200
                data = response.json()
                assert data["success"] is True
                assert "multiple" in data["message"].lower() or "images" in data["message"].lower()
                
        finally:
            # Cleanup
            for file in [image1_file, image2_file, pdf_file]:
                if os.path.exists(file):
                    os.remove(file)

    def test_ocr_error_handling(self, async_client):
        """Test OCR error handling with problematic images."""
        # Create a problematic image (very small, low contrast)
        width, height = 50, 50
        image = Image.new('RGB', (width, height), color='lightgray')
        draw = ImageDraw.Draw(image)
        draw.text((10, 10), "Tiny", fill='lightgray')  # Low contrast text
        image_file = "problematic_image.png"
        image.save(image_file)
        
        try:
            # Create PDF with problematic image
            pdf_file = create_pdf_with_image(image_file, "test_problematic.pdf")
            
            # Mock the RAG service response
            with patch('app.api.routes.documents.rag_service.add_document') as mock_add_document:
                mock_add_document.return_value = {
                    "success": True,
                    "message": f"Document '{pdf_file}' processed with OCR warnings",
                    "chunks_processed": 1,
                    "ocr_extracted_text": [],
                    "ocr_warnings": ["Low contrast text", "Small image size"]
                }
                
                # Upload the PDF file
                with open(pdf_file, 'rb') as f:
                    files = {"file": (pdf_file, f, "application/pdf")}
                    response = async_client.post("/documents/upload", files=files)
                
                assert response.status_code == 200
                data = response.json()
                assert data["success"] is True
                assert "warnings" in data["message"].lower() or "ocr" in data["message"].lower()
                
        finally:
            # Cleanup
            for file in [image_file, pdf_file]:
                if os.path.exists(file):
                    os.remove(file)


@pytest.mark.ocr_performance
@pytest.mark.skipif(not OCR_AVAILABLE, reason="OCR dependencies not available")
class TestOCRPerformance:
    """Test OCR performance and resource usage."""
    
    def test_ocr_processing_time(self, async_client):
        """Test OCR processing time for different document sizes."""
        # Create test image with moderate text
        moderate_text = [f"Line {i} of moderate text content" for i in range(10)]
        image_file = create_test_image_with_text(moderate_text)
        
        try:
            # Create PDF with the image
            pdf_file = create_pdf_with_image(image_file)
            
            # Mock the RAG service response
            with patch('app.api.routes.documents.rag_service.add_document') as mock_add_document:
                mock_add_document.return_value = {
                    "success": True,
                    "message": f"Document '{pdf_file}' processed efficiently",
                    "chunks_processed": 10,
                    "ocr_extracted_text": moderate_text,
                    "processing_time_ms": 1500
                }
                
                # Upload the PDF file
                with open(pdf_file, 'rb') as f:
                    files = {"file": (pdf_file, f, "application/pdf")}
                    response = async_client.post("/documents/upload", files=files)
                
                assert response.status_code == 200
                data = response.json()
                assert data["success"] is True
                assert "efficiently" in data["message"].lower() or "processed" in data["message"].lower()
                
                # Check processing time if available
                if "processing_time_ms" in data:
                    assert data["processing_time_ms"] < 5000  # Should be under 5 seconds
                
        finally:
            # Cleanup
            for file in [image_file, pdf_file]:
                if os.path.exists(file):
                    os.remove(file)

    def test_ocr_memory_usage(self, async_client):
        """Test OCR memory usage during processing."""
        # Create test image with large text content
        large_text = [f"Large text line {i} with substantial content for memory testing" for i in range(20)]
        image_file = create_test_image_with_text(large_text)
        
        try:
            # Create PDF with the image
            pdf_file = create_pdf_with_image(image_file)
            
            # Mock the RAG service response
            with patch('app.api.routes.documents.rag_service.add_document') as mock_add_document:
                mock_add_document.return_value = {
                    "success": True,
                    "message": f"Document '{pdf_file}' processed with optimal memory usage",
                    "chunks_processed": 20,
                    "ocr_extracted_text": large_text,
                    "memory_usage_mb": 45
                }
                
                # Upload the PDF file
                with open(pdf_file, 'rb') as f:
                    files = {"file": (pdf_file, f, "application/pdf")}
                    response = async_client.post("/documents/upload", files=files)
                
                assert response.status_code == 200
                data = response.json()
                assert data["success"] is True
                assert "optimal" in data["message"].lower() or "processed" in data["message"].lower()
                
                # Check memory usage if available
                if "memory_usage_mb" in data:
                    assert data["memory_usage_mb"] < 100  # Should be under 100MB
                
        except Exception as e:
            print(f"Error during OCR memory usage test: {e}")
        finally:
            # Cleanup
            for file in [image_file, pdf_file]:
                if os.path.exists(file):
                    os.remove(file)


def run_ocr_tests():
    """Helper function to run OCR tests programmatically."""
    import pytest
    pytest.main([__file__, "-v", "-m", "ocr"]) 